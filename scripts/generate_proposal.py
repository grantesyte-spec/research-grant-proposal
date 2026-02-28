#!/usr/bin/env python3
"""
Generate a DOCX grant proposal from markdown file.

Usage:
  python generate_proposal.py --title "课题名称" --markdown proposal.md --output proposal.docx
"""

import argparse
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def apply_doc_style(doc):
    """设置文档格式"""
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_title(doc, title):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()


def parse_markdown_file(markdown_path):
    """解析markdown文件"""
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    result = []
    table_data = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 标题
        if line.startswith('# '):
            result.append(('title', line[2:]))
        elif line.startswith('## '):
            result.append(('h2', line[3:]))
        elif line.startswith('### '):
            result.append(('h3', line[4:]))
        elif line.startswith('#### '):
            result.append(('h4', line[5:]))
        # 表格
        elif line.startswith('|') and not line.startswith('|---'):
            table_data.append(line)
        elif line.startswith('|---'):
            pass
        elif table_data and not line.startswith('|'):
            result.append(('table', table_data.copy()))
            table_data = []
            if line.startswith('|'):
                table_data.append(line)
        # 无序列表
        elif line.startswith('- ') or line.startswith('* '):
            result.append(('list', '  • ' + line[2:]))
        # 数字列表
        elif re.match(r'^\d+\. ', line):
            result.append(('numlist', '  ' + line))
        # 参考文献
        elif re.match(r'^\[\d+\]', line):
            result.append(('ref', line))
        # 普通段落
        else:
            result.append(('para', line))

    if table_data:
        result.append(('table', table_data))
    return result


def add_table(doc, table_lines):
    """添加表格"""
    if not table_lines:
        return

    rows = []
    for line in table_lines:
        if line.startswith('|---'):
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)

    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
    doc.add_paragraph()


def build_docx(title, markdown_path, output):
    """从markdown生成Word"""
    doc = Document()
    apply_doc_style(doc)
    add_title(doc, title)

    elements = parse_markdown_file(markdown_path)

    for elem_type, elem_content in elements:
        if elem_type == 'title':
            pass
        elif elem_type == 'h2':
            doc.add_paragraph()
            p = doc.add_paragraph(elem_content)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
        elif elem_type == 'h3':
            p = doc.add_paragraph(elem_content)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        elif elem_type == 'h4':
            p = doc.add_paragraph(elem_content)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif elem_type == 'para':
            p = doc.add_paragraph(elem_content)
            p.paragraph_format.first_line_indent = Inches(0.3)
        elif elem_type == 'list':
            p = doc.add_paragraph(elem_content)
        elif elem_type == 'numlist':
            p = doc.add_paragraph(elem_content)
        elif elem_type == 'ref':
            p = doc.add_paragraph(elem_content)
        elif elem_type == 'table':
            add_table(doc, elem_content)

    output = os.path.expanduser(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)
    doc.save(output)
    return output


def main():
    parser = argparse.ArgumentParser(description="Generate grant proposal DOCX from markdown")
    parser.add_argument("--title", required=True)
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--output", default="~/Desktop/proposal.docx")
    args = parser.parse_args()

    out = build_docx(args.title, args.markdown, args.output)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
